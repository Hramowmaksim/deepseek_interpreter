import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime

def parse_conversations(mapping):
    """Извлекает сообщения из древовидной структуры mapping"""
    messages = []
    current_node_id = "root"
    
    while True:
        current_node = mapping.get(current_node_id)
        if not current_node or not current_node.get("children"):
            break
            
        next_node_id = current_node["children"][0]
        next_node = mapping.get(next_node_id)
        
        if next_node and next_node.get("message"):
            messages.append({
                "role": "assistant" if next_node["message"].get("model") else "user",
                "content": next_node["message"]["content"],
                "time": next_node["message"]["inserted_at"]
            })
        
        current_node_id = next_node_id
    
    return messages

def create_chat_docx(chat_data, filename):
    doc = Document()
    
    # Настройка стиля документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Добавление заголовка чата
    title = chat_data.get('title', 'Без названия')
    create_time = chat_data.get('inserted_at', '')
    try:
        dt = datetime.fromisoformat(create_time)
        create_time = dt.strftime("%Y-%m-%d %H:%M:%S")
    except:
        pass
    
    doc.add_heading(f"{title} ({create_time})", level=1)
    
    # Извлечение сообщений
    messages = parse_conversations(chat_data.get('mapping', {}))
    
    # Добавление сообщений в документ
    for msg in messages:
        if msg["role"] == "user":
            # Сообщение пользователя
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = p.add_run(msg["content"])
            run.bold = True
        else:
            # Ответ DeepSeek
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.add_run(msg["content"])
        
        # Добавляем время сообщения
        time_paragraph = doc.add_paragraph()
        time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        time_run = time_paragraph.add_run(msg.get("time", ""))
        time_run.font.size = Pt(9)
        time_run.font.italic = True
        
        doc.add_paragraph()  # Пустая строка между сообщениями
    
    # Сохранение документа
    doc.save(filename)

def main():
    # Проверка существования файла
    if not os.path.exists('conversations.json'):
        print("Ошибка: файл conversations.json не найден в текущей директории.")
        return
    
    # Чтение JSON файла
    try:
        with open('conversations.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        return
    
    # Если файл содержит один чат (как в примере)
    if isinstance(data, dict):
        data = [data]
    
    # Создание папки для чатов, если её нет
    if not os.path.exists('chats'):
        os.makedirs('chats')
    
    # Обработка каждого чата
    for i, chat in enumerate(data, start=1):
        chat_title = chat.get('title', f'chat_{i}')
        # Заменяем недопустимые символы в имени файла
        safe_title = "".join(c if c.isalnum() or c in " _-()" else "_" for c in chat_title)
        filename = f"chats/{safe_title}.docx"
        create_chat_docx(chat, filename)
        print(f"Создан файл: {filename}")
    
    print("Готово! Все чаты сохранены в папке 'chats'.")

if __name__ == "__main__":
    main()